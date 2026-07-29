from __future__ import annotations

import textwrap
import zipfile
from datetime import datetime
from pathlib import Path


ORIGIN_TEXT = (
    "Trace File Lineage deterministic PDF integration fixture. "
    "The editable source describes cedar samples, cobalt markers, and reproducible archive evidence. "
    "A local exporter converts this document into a text-bearing PDF without network services. "
    "The provenance scanner should recover this distinctive wording and rank the editable document first. "
    "Embedded media carries the same blue and green reference panel in both containers. "
    "No filename stem is shared, so content and media evidence must drive the origin ranking."
)

UNRELATED_TEXT = (
    "Unrelated meeting record about budgets, travel dates, catering, and room reservations. "
    "This content intentionally shares no distinctive provenance phrases with the PDF fixture."
)

OCR_TEXT = "LINEAGE OCR RUNTIME 7319"
APPENDIX_TEXT = (
    "Independent appendix source for the multi-source export. "
    "It contributes quartz calibration notes and the identifier APPENDIX-8421."
)


def _normalize_docx_zip(path: Path) -> None:
    temporary = path.with_suffix(".normalized.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for name in sorted(source.namelist()):
            original = source.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(2020, 1, 2, 3, 4, 6))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            target.writestr(info, source.read(name))
    temporary.replace(path)


def _write_reference_image(path: Path, *, alternate: bool = False) -> None:
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (360, 160), "white")
    draw = ImageDraw.Draw(image)
    if alternate:
        draw.rectangle((20, 20, 340, 140), fill=(174, 71, 86), outline=(71, 23, 32), width=5)
        draw.text((78, 68), "UNRELATED MEDIA", fill="white")
    else:
        draw.rectangle((20, 20, 170, 140), fill=(31, 111, 178), outline=(14, 55, 92), width=5)
        draw.rectangle((190, 20, 340, 140), fill=(54, 145, 92), outline=(24, 78, 48), width=5)
        draw.text((92, 68), "MEDIA 7319", fill="white")
    image.save(path, format="PNG", optimize=False, compress_level=9)


def _write_ocr_image(path: Path) -> None:
    from PIL import Image, ImageDraw, ImageFont

    image = Image.new("L", (1400, 320), 255)
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 72)
    except OSError:
        font = ImageFont.load_default(size=72)
    draw.text((60, 110), OCR_TEXT, font=font, fill=0, stroke_width=1)
    image.save(path, format="PNG", optimize=False, compress_level=9)


def _write_docx(path: Path, text: str, image_path: Path, title: str) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)
    document.add_heading(title, level=1)
    for paragraph in text.split(". "):
        document.add_paragraph(paragraph.rstrip(".") + ".")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Fixture"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Marker"
    table.cell(1, 1).text = "7319"
    document.add_picture(str(image_path), width=Inches(3.6))
    properties = document.core_properties
    properties.author = "Trace File Lineage"
    properties.title = title
    properties.created = datetime(2020, 1, 2, 3, 4, 6)
    properties.modified = datetime(2020, 1, 2, 3, 4, 6)
    document.save(path)
    _normalize_docx_zip(path)


def _write_text_pdf(path: Path, image_path: Path, text_value: str = ORIGIN_TEXT, title_value: str = "Deterministic PDF Integration Fixture") -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=letter, pageCompression=0, invariant=1)
    pdf.setAuthor("Trace File Lineage")
    pdf.setTitle("Deterministic text-bearing PDF fixture")
    text = pdf.beginText(72, 720)
    text.setFont("Helvetica-Bold", 16)
    text.textLine(title_value)
    text.moveCursor(0, 14)
    text.setFont("Helvetica", 10)
    for line in textwrap.wrap(text_value, width=88):
        text.textLine(line)
    pdf.drawText(text)
    pdf.drawImage(ImageReader(str(image_path)), 72, 315, width=360, height=160, mask="auto")
    pdf.showPage()
    pdf.save()


def _write_scanned_pdf(path: Path, image_path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=letter, pageCompression=1, invariant=1)
    pdf.setTitle("Image-only OCR fixture")
    pdf.drawImage(ImageReader(str(image_path)), 36, 330, width=540, height=123, mask="auto")
    pdf.showPage()
    pdf.save()


def generate(output_dir: Path) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    media = output_dir / "reference-media.png"
    unrelated_media = output_dir / "unrelated-media.png"
    ocr_png = output_dir / "scanned-text.png"
    _write_reference_image(media)
    _write_reference_image(unrelated_media, alternate=True)
    _write_ocr_image(ocr_png)
    source = output_dir / "proposal_master.docx"
    unrelated = output_dir / "meeting_notes.docx"
    old_source = output_dir / "proposal_old_revision.docx"
    appendix = output_dir / "appendix_source.docx"
    pdf = output_dir / "submission_final.pdf"
    post_edited = output_dir / "submission_post_edited.pdf"
    multi_source = output_dir / "submission_multi_source.pdf"
    scanned_pdf = output_dir / "scanned-text.pdf"
    _write_docx(source, ORIGIN_TEXT, media, "Editable provenance source")
    _write_docx(unrelated, UNRELATED_TEXT, unrelated_media, "Unrelated meeting notes")
    _write_docx(old_source, ORIGIN_TEXT[:230] + " Obsolete revision marker OLD-110.", media, "Old editable revision")
    _write_docx(appendix, APPENDIX_TEXT, unrelated_media, "Independent appendix")
    _write_text_pdf(pdf, media)
    _write_text_pdf(post_edited, media, ORIGIN_TEXT + " A manual post-export edit added REVIEWED-2026.", "Post-edited PDF fixture")
    _write_text_pdf(multi_source, media, ORIGIN_TEXT + " " + APPENDIX_TEXT, "Multi-source PDF fixture")
    _write_scanned_pdf(scanned_pdf, ocr_png)
    return {
        "source_docx": source,
        "unrelated_docx": unrelated,
        "old_source_docx": old_source,
        "appendix_docx": appendix,
        "text_pdf": pdf,
        "post_edited_pdf": post_edited,
        "multi_source_pdf": multi_source,
        "media_png": media,
        "ocr_png": ocr_png,
        "scanned_pdf": scanned_pdf,
    }


if __name__ == "__main__":
    import argparse
    import json

    parser = argparse.ArgumentParser()
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    print(json.dumps({key: str(value) for key, value in generate(args.output_dir).items()}, indent=2, sort_keys=True))
